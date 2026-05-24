from docx import Document
from config import EXPORT_DOCX_DIR
from services.export_utils import safe_course_filename


def export_docx(task, slides):
    doc = Document()
    doc.add_heading(task["course_title"], level=1)
    summary = doc.add_paragraph()
    summary.add_run("Grade: ").bold = True
    summary.add_run(f"{task.get('grade', '')}    ")
    summary.add_run("Textbook: ").bold = True
    summary.add_run(f"{task.get('textbook', '')}    ")
    summary.add_run("Unit: ").bold = True
    summary.add_run(f"{task.get('unit', '')}    ")
    summary.add_run("Lesson Type: ").bold = True
    summary.add_run(f"{task.get('lesson_type', '')}")

    doc.add_heading("Teacher Script", level=2)
    for slide in slides:
        doc.add_heading(f"Slide {slide['slide_index']}: {slide['title']}", level=3)
        doc.add_paragraph("Visible Content:")
        for item in slide.get("visible_content", []):
            doc.add_paragraph(str(item), style="List Bullet")
        if slide.get("key_sentence"):
            doc.add_paragraph(f"Key Sentence: {slide.get('key_sentence', '')}")
        if slide.get("useful_expressions"):
            doc.add_paragraph("Useful Expressions:")
            for item in slide.get("useful_expressions", []):
                doc.add_paragraph(str(item), style="List Bullet")
        if slide.get("possible_answers"):
            doc.add_paragraph("Possible Answers:")
            for item in slide.get("possible_answers", []):
                doc.add_paragraph(str(item), style="List Bullet")
        if slide.get("image_suggestion"):
            doc.add_paragraph(f"Image Suggestion: {slide.get('image_suggestion', '')}")
        if slide.get("chinese_hint"):
            doc.add_paragraph(f"Chinese Hint: {slide.get('chinese_hint', '')}")
        doc.add_paragraph(f"Teacher Notes: {slide.get('teacher_notes', '')}")
        doc.add_paragraph(f"Teaching Purpose: {slide.get('teaching_purpose', '')}")
        doc.add_paragraph(f"Estimated Time: {slide.get('estimated_time', '')}")
        doc.add_paragraph(f"Interaction Type: {slide.get('interaction_type', '')}")

    filename = safe_course_filename(task, "script.docx")
    path = EXPORT_DOCX_DIR / filename
    doc.save(path)
    return path
