"""Document parsing helpers shared by manuscript and knowledge workflows.

This module is deterministic and local-only. It does not call any LLM.
"""

from pathlib import Path
import re

from docx import Document


SUPPORTED_MANUSCRIPT_EXTENSIONS = {".txt", ".md", ".docx", ".pdf"}


class DocumentParseError(Exception):
    """Raised when a document cannot be parsed into usable text."""


def clean_extracted_text(text):
    """Normalize whitespace while preserving paragraph structure."""

    normalized = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in normalized.split("\n")]
    cleaned_lines = []
    blank_run = 0
    for line in lines:
        if not line:
            blank_run += 1
            if blank_run <= 2:
                cleaned_lines.append("")
            continue
        blank_run = 0
        cleaned_lines.append(line)
    cleaned = "\n".join(cleaned_lines).strip()
    return re.sub(r"\n{3,}", "\n\n", cleaned)


def _read_text_with_fallbacks(file_path):
    path = Path(file_path)
    errors = []
    for encoding in ("utf-8", "gbk"):
        try:
            return path.read_text(encoding=encoding)
        except (UnicodeDecodeError, LookupError) as exc:
            errors.append(str(exc))
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except OSError as exc:
        raise DocumentParseError(f"文本文件读取失败：{exc}") from exc
    except Exception as exc:
        detail = " | ".join(errors[:2]) if errors else str(exc)
        raise DocumentParseError(f"文本文件编码无法识别：{detail}") from exc


def extract_text_from_txt(file_path):
    return _read_text_with_fallbacks(file_path)


def extract_text_from_markdown(file_path):
    return _read_text_with_fallbacks(file_path)


def extract_text_from_docx(file_path):
    try:
        document = Document(file_path)
    except Exception as exc:
        raise DocumentParseError(f"DOCX 解析失败：{exc}") from exc

    chunks = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                chunks.append("\t".join(cells))

    return "\n".join(chunks)


def extract_text_from_pdf(file_path):
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentParseError("PDF 解析依赖 pypdf，当前环境未安装。") from exc

    try:
        reader = PdfReader(file_path)
    except Exception as exc:
        raise DocumentParseError(f"PDF 打开失败：{exc}") from exc

    chunks = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            chunks.append(text.strip())

    if not chunks:
        raise DocumentParseError("PDF 可能是扫描件，当前版本暂不支持 OCR。")
    return "\n\n".join(chunks)


def extract_text_from_file(file_path, filename=None):
    """Extract plain text from txt/md/docx/pdf files."""

    path = Path(file_path)
    suffix = (path.suffix or Path(filename or "").suffix).lower()
    if suffix not in SUPPORTED_MANUSCRIPT_EXTENSIONS:
        raise DocumentParseError(f"不支持的文件格式：{suffix or 'unknown'}")

    if suffix == ".txt":
        text = extract_text_from_txt(path)
    elif suffix == ".md":
        text = extract_text_from_markdown(path)
    elif suffix == ".docx":
        text = extract_text_from_docx(path)
    elif suffix == ".pdf":
        text = extract_text_from_pdf(path)
    else:
        raise DocumentParseError(f"不支持的文件格式：{suffix}")

    cleaned = clean_extracted_text(text)
    if not cleaned:
        raise DocumentParseError("文案解析后内容为空。")
    return cleaned


def save_parsed_text(text, target_dir, base_name):
    target = Path(target_dir)
    target.mkdir(parents=True, exist_ok=True)
    safe_base = re.sub(r"[^A-Za-z0-9._-]+", "_", str(base_name or "knowledge_text")).strip("._") or "knowledge_text"
    file_path = target / f"{safe_base}.txt"
    file_path.write_text(str(text or ""), encoding="utf-8")
    return str(file_path)


def count_text_words(text):
    """Rough mixed-language count.

    English is counted by word; Chinese is counted by Han character.
    """

    content = str(text or "")
    chinese_chars = re.findall(r"[\u4e00-\u9fff]", content)
    english_words = re.findall(r"[A-Za-z0-9']+", content)
    return len(chinese_chars) + len(english_words)


def count_source_words(text):
    """Backward-compatible alias used by manuscript workflow."""

    return count_text_words(text)


def generate_basic_summary(text, title=None, doc_type=None):
    body = clean_extracted_text(text)
    if not body:
        return ""

    prefix_parts = []
    if title:
        prefix_parts.append(str(title).strip())
    if doc_type:
        prefix_parts.append(str(doc_type).strip())
    prefix = "｜".join(part for part in prefix_parts if part)
    snippet = body[:300]
    if len(body) <= 300:
        return f"{prefix}：{snippet}" if prefix else snippet
    return f"{prefix}：{snippet}..." if prefix else f"{snippet}..."
