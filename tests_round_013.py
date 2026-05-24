import io
import re
import sqlite3
import uuid
from pathlib import Path

from docx import Document

from app import create_app
from config import DATABASE_PATH


def _extract_doc_id(location):
    match = re.search(r"/knowledge/(\d+)", str(location or ""))
    return int(match.group(1)) if match else None


def _build_docx_bytes():
    buffer = io.BytesIO()
    document = Document()
    document.add_paragraph("Lesson Plan: The Power of Reading")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Objective"
    table.rows[0].cells[1].text = "Improve reading comprehension."
    document.save(buffer)
    buffer.seek(0)
    return buffer


def run():
    app = create_app()
    app.config["TESTING"] = True
    title_prefix = f"[TEST] Round13 {uuid.uuid4().hex[:8]}"
    created_doc_ids = []

    with sqlite3.connect(DATABASE_PATH) as conn:
        exists = conn.execute(
            "SELECT name FROM sqlite_master WHERE type='table' AND name='knowledge_documents'"
        ).fetchone()
        assert exists, "knowledge_documents table was not created."

    with app.test_client() as client:
        # Old feature pages remain accessible.
        assert client.get("/").status_code == 200
        assert client.get("/ppt/new").status_code == 200
        assert client.get("/ppt/from-manuscript").status_code == 200
        assert client.get("/ppt/tasks").status_code == 200
        assert client.get("/settings/ai-models").status_code == 200
        assert client.get("/settings/agent-bindings").status_code == 200

        # Knowledge pages.
        assert client.get("/knowledge").status_code == 200
        assert client.get("/knowledge/new").status_code == 200

        # Raw text create.
        resp = client.post(
            "/knowledge/new",
            data={
                "title": f"{title_prefix} Raw",
                "doc_type": "教案",
                "grade": "高一",
                "textbook": "人教版",
                "volume": "必修一",
                "unit": "Unit 1",
                "lesson_type": "Reading",
                "tags": "test,raw",
                "raw_text": "This is a reading lesson draft.\nStudents will predict and discuss.",
            },
            follow_redirects=False,
        )
        assert resp.status_code in (302, 303), "Raw text create did not redirect."
        doc_id = _extract_doc_id(resp.headers.get("Location"))
        assert doc_id, "Cannot parse new knowledge doc id from redirect."
        created_doc_ids.append(doc_id)

        detail = client.get(f"/knowledge/{doc_id}")
        assert detail.status_code == 200
        assert "reading lesson draft".lower() in detail.get_data(as_text=True).lower()

        reparsed = client.post(f"/knowledge/{doc_id}/reparse", follow_redirects=True)
        assert reparsed.status_code == 200
        text_resp = client.get(f"/knowledge/{doc_id}/text")
        assert text_resp.status_code == 200

        # TXT upload.
        txt_resp = client.post(
            "/knowledge/new",
            data={
                "title": f"{title_prefix} TXT",
                "doc_type": "阅读材料",
                "grade": "高二",
                "textbook": "外研版",
                "volume": "必修二",
                "unit": "Unit 2",
                "lesson_type": "Reading",
                "tags": "test,txt",
                "raw_text": "",
                "file": (io.BytesIO("Line one.\nLine two.".encode("utf-8")), "sample.txt"),
            },
            content_type="multipart/form-data",
            follow_redirects=False,
        )
        assert txt_resp.status_code in (302, 303)
        txt_id = _extract_doc_id(txt_resp.headers.get("Location"))
        assert txt_id
        created_doc_ids.append(txt_id)

        # MD upload.
        md_resp = client.post(
            "/knowledge/new",
            data={
                "title": f"{title_prefix} MD",
                "doc_type": "PPT文案",
                "grade": "高二",
                "textbook": "人教版",
                "volume": "必修三",
                "unit": "Unit 3",
                "lesson_type": "Writing",
                "tags": "test,md",
                "raw_text": "",
                "file": (io.BytesIO("# Title\n- point 1\n- point 2\n".encode("utf-8")), "sample.md"),
            },
            content_type="multipart/form-data",
            follow_redirects=False,
        )
        assert md_resp.status_code in (302, 303)
        md_id = _extract_doc_id(md_resp.headers.get("Location"))
        assert md_id
        created_doc_ids.append(md_id)

        # DOCX upload.
        docx_resp = client.post(
            "/knowledge/new",
            data={
                "title": f"{title_prefix} DOCX",
                "doc_type": "教案",
                "grade": "高三",
                "textbook": "译林版",
                "volume": "选择性必修一",
                "unit": "Unit 4",
                "lesson_type": "Grammar",
                "tags": "test,docx",
                "raw_text": "",
                "file": (_build_docx_bytes(), "sample.docx"),
            },
            content_type="multipart/form-data",
            follow_redirects=False,
        )
        assert docx_resp.status_code in (302, 303)
        docx_id = _extract_doc_id(docx_resp.headers.get("Location"))
        assert docx_id
        created_doc_ids.append(docx_id)

        # List query and keyword filter.
        list_resp = client.get(f"/knowledge?keyword={title_prefix}")
        assert list_resp.status_code == 200
        assert title_prefix in list_resp.get_data(as_text=True)

        with sqlite3.connect(DATABASE_PATH) as conn:
            conn.row_factory = sqlite3.Row
            row = conn.execute(
                "SELECT * FROM knowledge_documents WHERE id=?",
                (doc_id,),
            ).fetchone()
            assert row, "Knowledge row not found."
            assert int(row["word_count"] or 0) > 0, "word_count should be > 0."
            assert str(row["summary"] or "").strip(), "summary should not be empty."

        # Delete flow.
        del_resp = client.post(f"/knowledge/{docx_id}/delete", follow_redirects=False)
        assert del_resp.status_code in (302, 303)
        created_doc_ids.remove(docx_id)

    # Cleanup remaining created records.
    with sqlite3.connect(DATABASE_PATH) as conn:
        conn.execute(
            "DELETE FROM knowledge_documents WHERE title LIKE ?",
            (f"{title_prefix}%",),
        )
        conn.commit()

    print("ROUND_013_OK")


if __name__ == "__main__":
    run()
